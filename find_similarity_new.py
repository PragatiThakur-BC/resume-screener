import streamlit as st
from gensim.models.doc2vec import Doc2Vec
from numpy.linalg import norm
import numpy as np
import re
from PyPDF2 import PdfReader
import os
import pandas as pd
import base64
from io import BytesIO
from docx import Document

# Load the trained model
model = Doc2Vec.load('cv_job_matching_dsp.model')

# Function to preprocess text
def preprocess_text(text):
    text = text.lower()
    text = re.sub('[^a-z]', ' ', text)
    text = re.sub(r'\d+', '', text)
    text = ' '.join(text.split())
    return text

# Function to calculate similarity score between resume and job description
def calculate_match_score(resume_text, jd_text):

    # Preprocess text
    preprocessed_resume = preprocess_text(resume_text)
    preprocessed_jd = preprocess_text(jd_text)

    # Get embeddings
    v1 = model.infer_vector(preprocessed_resume.split())
    v2 = model.infer_vector(preprocessed_jd.split())

    # Calculate similarity
    similarity = 100 * (np.dot(np.array(v1), np.array(v2))) / (norm(np.array(v1)) * norm(np.array(v2)))
    return similarity

# Function to select job description PDF file
def select_job_description():
    st.header("Job Description")
    jd_text = st.text_area("Paste your job description here:", height=200)
    return jd_text

# Function to select multiple resume files
def select_resumes():
    st.header("Select Resumes")
    uploaded_files = st.file_uploader("Upload up to 50 resume files", accept_multiple_files=True)
    return uploaded_files

def main():
    st.title("Blenheim Chalcot Resume Screener")
    jd_text = select_job_description()
    uploaded_files = select_resumes()

    if st.button("Screen Resumes"):
        if jd_text is None or jd_text == "":
            st.error("Please paste a Job Description.")
        elif uploaded_files is None or len(uploaded_files) == 0:
            st.error("Please upload at least one Resume file.")
        else:
            # Clear main body content
            st.empty()
            # Display the ranking table
            display_ranking_table(uploaded_files, jd_text)

# Function to display the ranking table
def display_ranking_table(uploaded_files, jd_text):

    resume_ranking = []
    for rank, resume_file in enumerate(uploaded_files, start=1):
        try:
            # Extract text from the file
            if resume_file.name.endswith(".pdf"):
                resume_text = extract_text_from_pdf(resume_file)
            elif resume_file.name.endswith((".txt", ".docx")):
                resume_text = extract_text_from_other_file(resume_file)
            else:
                st.warning(f"Unsupported file type: {resume_file.name}")
                continue

            # Calculate similarity score
            similarity_score = round(calculate_match_score(resume_text, jd_text), 2)

            reason = ""
            if similarity_score > 80:
                reason = f"Resume '{resume_file.name}' received a high similarity score of {similarity_score:.2f} because it covered most of the key skills and responsibilities as per the job description."
            elif similarity_score > 65 and similarity_score <= 80:
                reason = f"Resume '{resume_file.name}' received a good similarity score of {similarity_score:.2f} because it discussed experiences and skills that were relevant to the job role."
            elif similarity_score > 50 and similarity_score <= 65:
                reason = f"Resume '{resume_file.name}' received a moderate similarity score of {similarity_score:.2f}. There is a partial keyword overlap with the job description."
            elif similarity_score > 35 and similarity_score <= 50:
                reason = f"Resume '{resume_file.name}' received a low similarity score of {similarity_score:.2f}. There is limited keyword overlap with the job description."
            else:
                reason = f"Resume '{resume_file.name}' received a very low similarity score of {similarity_score:.2f}. There is minimal keyword overlap with the job description."

            # Append to resume ranking list
            resume_ranking.append((rank, resume_file.name, similarity_score, reason))
        except Exception as e:
            st.write(f"Not a proper file format,Recheck the file : {resume_file.name}: {e}")

    # Sort the resume ranking based on similarity score
    resume_ranking.sort(key=lambda x: x[2], reverse=True)

    # Update rank column based on the sorted order
    for idx, (rank, resume_name, similarity_score, reason) in enumerate(resume_ranking, start=1):
        resume_ranking[idx - 1] = (idx, resume_name, similarity_score, reason)

    # Create DataFrame without index column
    df = pd.DataFrame(resume_ranking, columns=["Rank", "Resume Name", "Similarity Score(%)", "Reason"])

    # Set "Rank" column as index
    df.set_index("Rank", inplace=True)

    # Format similarity score column to display two decimal places
    df["Similarity Score(%)"] = df["Similarity Score(%)"].apply(lambda x: f"{x:.2f}")

    # Display the ranking table
    st.markdown("<h2>Resume Ranking</h2>", unsafe_allow_html=True)
    st.table(df)

    # Generate and download Excel sheet
    generate_excel(df)

# Function to extract text from PDF file
def extract_text_from_pdf(pdf_file):
    pdf_text = ""
    with pdf_file as file:
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            pdf_text += page.extract_text()
    return pdf_text

# Function to extract text from TXT or DOCX file
def extract_text_from_other_file(file):
    file_content = file.read()
    if file.name.endswith(".txt"):
        resume_text = file_content.decode("utf-8")
    elif file.name.endswith(".docx"):
        doc = Document(BytesIO(file_content))
        resume_text = ""
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            resume_text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    resume_text += cell.text + "\n"

    return resume_text.strip()

# Function to generate Excel sheet
def generate_excel(df):
    # Add Rank column to DataFrame for Excel
    df_excel = df.copy()
    df_excel.insert(0, "Rank", df.index)

    # Convert DataFrame to Excel file
    excel_file = save_excel_to_buffer(df_excel)

    # Download Excel file
    st.markdown("<p><a href='{0}' download='resume_ranking.xlsx'>Download Excel file</a></p>".format(get_excel_download_link(excel_file)), unsafe_allow_html=True)

# Function to save DataFrame to Excel file in memory buffer
def save_excel_to_buffer(df):
    excel_buffer = BytesIO()
    df.to_excel(excel_buffer, index=False)
    excel_buffer.seek(0)
    return excel_buffer

# Function to create a download link for the Excel file
def get_excel_download_link(excel_buffer):
    excel_content = excel_buffer.getvalue()
    b64 = base64.b64encode(excel_content).decode()
    href = f'data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}'
    return href

# Image in the sidebar with responsive width
st.sidebar.image("images2.ico", use_column_width=True)

st.sidebar.markdown(
    """    
    ## Welcome to the Blenheim Chalcot Resume Screener!

    This application is designed to streamline the resume screening process for the HR team at BlenheimChalcot.
    By leveraging advanced technology, it empowers you to make more informed decisions efficiently.

    **How it works:**
    - Upload up to 50 resumes and the job description.
    - Our intelligent screening algorithm will analyze the resumes based on the uploaded job description.
    - You'll receive instant feedback on candidate suitability, saving you time and effort.

    **Key Features of Resume Screener:**
    - Faster screening: Say goodbye to manual resume sorting.
    - Improved accuracy: Our algorithm identifies the best-fit candidates for your roles.
    - Enhanced efficiency: Spend less time on administrative tasks and more time engaging with top talent.
    """
)

# Adding CSS styles
st.markdown(
    """
    <style> 
        body { 
            background-color: #bcbcbc; 
        } 
        .st-bw {  
            padding: 20px; 
            border-radius: 10px; 
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1); 
        } 
        .stButton>button { 
            width: 200px !important; 
            height: 50px !important; 
            font-size: 20px !important; 
            background-color: #ff0066 !important; 
            color: white !important; 
            border-radius: 5px !important; 
            border: none !important; 
            cursor: pointer !important; 
            transition: all 0.3s ease !important; 
            margin-top: 20px !important; 
        } 
        .stButton>button:hover { 
            background-color: #f74557 !important; 
        } 
    </style>
    """,
    unsafe_allow_html=True
)


# Run the main function
if __name__ == "__main__":
    main()
